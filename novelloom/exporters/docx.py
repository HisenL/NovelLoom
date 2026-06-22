from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from .base import ExportChapter, Exporter


class DocxExporter(Exporter):
    format = "docx"

    def export(self, *, title: str, chapters: list[ExportChapter], output: Path) -> Path:
        output.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        document.core_properties.title = title
        document.add_heading(title, level=0)
        normal = document.styles["Normal"]
        normal.font.name = "宋体"
        normal.font.size = Pt(11)
        for chapter in sorted(chapters, key=lambda item: item.order):
            document.add_heading(chapter.title, level=1)
            for block in chapter.markdown.split("\n\n"):
                text = block.strip()
                if text:
                    document.add_paragraph(text)
        document.save(str(output))
        return output
